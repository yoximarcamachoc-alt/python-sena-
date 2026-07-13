import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RUTA = r"C:\Users\SENA\Documents\VS_CODE\Sena_IA_1\IA_1\4-Preparación de datasets para aprendizaje no supervisado\G4V1_EST\05_informe"
NAVY = RGBColor(0x1B, 0x3A, 0x6B)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
GRAY = RGBColor(0xAA, 0xAA, 0xAA)

doc = Document()

# ---- Estilos ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = NAVY
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(16)
    elif level == 2:
        h.font.size = Pt(13)
    else:
        h.font.size = Pt(11)

def add_observation(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(10)
    return p

def set_cell_shading(cell, color="1B3A6B"):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        set_cell_shading(cell)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, "F2F6FA")
    return table

def add_figure(img_path, caption, width=5.5):
    full = os.path.join(RUTA, img_path)
    if os.path.exists(full):
        doc.add_picture(full, width=Inches(width))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
    else:
        doc.add_paragraph(f"[Imagen no encontrada: {img_path}]")

# ===== PORTADA =====
for _ in range(4):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Servicio Nacional de Aprendizaje (SENA)")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Centro de Gestión de Mercados, Logística y Tecnologías de la Información")
r.font.size = Pt(12)
r.font.color.rgb = BLACK

doc.add_paragraph()
doc.add_paragraph()

prog = doc.add_paragraph()
prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = prog.add_run("Programa: Análisis y Desarrollo de Inteligencia de Negocios")
r.font.size = Pt(12)

doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Guía 4 - Segmentación mediante Aprendizaje No Supervisado (K-Means)")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = NAVY

doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Informe Final Completo")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = NAVY

doc.add_paragraph()
for line in ["Instrucción: John Jairo Londoño", "Aprendiz: [Nombre del Aprendiz]", "Ficha: [Número de Ficha]", "Fecha: Junio 2026"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(11)

doc.add_page_break()

# ===== 1. INTRODUCCIÓN =====
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'El presente informe documenta el proceso completo de segmentación de clientes mediante la técnica '
    'de aprendizaje no supervisado K-Means, desarrollado en el marco de la Guía 4 del programa Análisis '
    'y Desarrollo de Inteligencia de Negocios del SENA. El objetivo principal es descubrir patrones ocultos '
    'en el comportamiento de los clientes a partir de un conjunto de variables numéricas y categóricas, '
    'sin utilizar una variable objetivo predefinida.'
)
doc.add_paragraph(
    'La segmentación de clientes es una técnica fundamental en inteligencia de negocios que permite '
    'agrupar individuos con características similares para diseñar estrategias comerciales, de retención '
    'y de marketing diferenciadas. En este estudio se empleó el algoritmo K-Means por su eficiencia '
    'computacional y facilidad de interpretación.'
)
doc.add_paragraph(
    'El análisis se realizó sobre una base de 420 clientes con 18 variables, de las cuales 15 fueron '
    'seleccionadas para el proceso de clustering. Se evaluaron valores de k desde 2 hasta 10 utilizando '
    'el método del codo (inercia), el coeficiente de Silhouette, y se complementó con visualizaciones '
    'de dispersión, PCA, diagramas de caja y matriz de correlaciones.'
)

# ===== 2. OBJETIVOS =====
doc.add_heading('2. Objetivos', level=1)
doc.add_heading('2.1 Objetivo General', level=2)
doc.add_paragraph(
    'Segmentar la base de clientes del dataset G4_base_clientes.csv utilizando el algoritmo K-Means '
    'para identificar grupos homogéneos de clientes que permitan diseñar estrategias de negocio '
    'diferenciadas.'
)
doc.add_heading('2.2 Objetivos Específicos', level=2)
objs = [
    'Realizar un diagnóstico inicial del dataset identificando tipos de datos, valores nulos, duplicados y variables disponibles.',
    'Seleccionar y preparar las variables relevantes para clustering mediante codificación (One-Hot y Ordinal) y escalamiento (StandardScaler).',
    'Determinar el número óptimo de clusters k mediante el método del codo y el coeficiente de Silhouette para k=2 hasta k=10.',
    'Entrenar el modelo K-Means con k=3 y asignar etiquetas de cluster a cada cliente.',
    'Visualizar los clusters mediante gráficas de dispersión, PCA y diagramas de caja.',
    'Caracterizar cada segmento mediante perfiles numéricos y categóricos.',
    'Realizar un análisis comparativo externo con la variable Abandono para validar la utilidad de los segmentos.'
]
for o in objs:
    doc.add_paragraph(o, style='List Bullet')

# ===== 3. METODOLOGÍA =====
doc.add_heading('3. Metodología', level=1)
doc.add_paragraph(
    'La metodología empleada sigue el flujo de trabajo estándar para proyectos de clustering con K-Means:'
)
steps = [
    'Carga y exploración inicial del dataset.',
    'Diagnóstico de calidad de datos (nulos, duplicados, tipos).',
    'Selección de variables: exclusión de identificadores y variables supervisadas (Abandono).',
    'Clasificación de variables en numéricas, nominales y ordinales.',
    'Preprocesamiento: escalamiento (StandardScaler) para numéricas, One-Hot Encoding para nominales, Ordinal Encoding para ordinales.',
    'Evaluación de k=2 a k=10 con método del codo (inercia) y Silhouette Score.',
    'Selección justificada de k=3.',
    'Entrenamiento del modelo final con k=3.',
    'Visualización: dispersión sin/con clusters, PCA, barras, boxplots, correlación.',
    'Perfilamiento numérico y categórico de cada cluster.',
    'Análisis externo con variable Abandono.'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ===== 4. DATASET =====
doc.add_heading('4. Dataset', level=1)
doc.add_heading('4.1 Descripción General', level=2)
doc.add_paragraph(
    'El dataset G4_base_clientes.csv contiene información de 420 clientes distribuidos en 18 columnas. '
    'No se encontraron valores nulos ni registros duplicados.'
)

headers = ['Variable', 'Tipo', 'Descripción']
rows = [
    ['ID_Cliente', 'str (identificador)', 'Código único del cliente'],
    ['Edad', 'int64', 'Edad del cliente en años'],
    ['IngresoMensual', 'int64', 'Ingreso mensual en unidades monetarias'],
    ['CantidadCompras', 'int64', 'Cantidad total de compras realizadas'],
    ['ComprasUltimos12M', 'int64', 'Compras en los últimos 12 meses'],
    ['AntiguedadMeses', 'int64', 'Antigüedad del cliente en meses'],
    ['QuejasUltimos6M', 'int64', 'Número de quejas en los últimos 6 meses'],
    ['DiasDesdeUltimaCompra', 'int64', 'Días desde la última compra'],
    ['VisitasWebUltimoMes', 'int64', 'Visitas al sitio web en el último mes'],
    ['TiempoPromedioSesionMin', 'float64', 'Tiempo promedio de sesión en minutos'],
    ['CuponesUsados', 'int64', 'Número de cupones utilizados'],
    ['Ciudad', 'str (nominal)', 'Ciudad de residencia'],
    ['CanalPreferido', 'str (nominal)', 'Canal de preferencia (Web, Tienda, Teléfono, App)'],
    ['ZonaResidencia', 'str (nominal)', 'Zona de residencia (Urbana, Suburbana, Rural)'],
    ['Segmento', 'str (ordinal)', 'Segmento del cliente (Básico, Medio, Premium)'],
    ['Satisfaccion', 'str (ordinal)', 'Nivel de satisfacción (Baja, Media, Alta)'],
    ['CodigoCampania', 'str', 'Código de campaña asignado'],
    ['Abandono', 'int64', 'Indicador de abandono (0=No, 1=Sí) - NO usado para clustering'],
]
make_table(headers, rows)

add_observation('Nota: La variable Abandono no se utilizó para entrenar el clustering. Solo se empleó al final como comparación externa para validar los segmentos encontrados.')

# ===== 5. SELECCIÓN DE k =====
doc.add_heading('5. Determinación del Número Óptimo de Clusters (k)', level=1)

doc.add_heading('5.1 Método del Codo', level=2)
doc.add_paragraph(
    'El método del codo evalúa la inercia (suma de distancias al cuadrado de cada punto a su centroide) '
    'para distintos valores de k. El "codo" es el punto donde la reducción de la inercia se vuelve marginal, '
    'indicando un equilibrio entre ajuste y simplicidad.'
)
doc.add_paragraph('Se evaluaron k desde 2 hasta 10, obteniendo los siguientes resultados:')

import pandas as pd
comp = pd.read_csv(os.path.join(RUTA.replace('05_informe', '04_notebook'), 'G4_comparacion_k_codo_silhouette_V2.csv'))
codo_headers = ['k', 'Inercia', 'Reducción Inercia', 'Reducción %', 'Silhouette Score']
codo_rows = []
for _, row in comp.iterrows():
    codo_rows.append([int(row['k']), f"{row['inercia']:.2f}",
                      f"{row['reduccion_inercia']:.2f}" if pd.notna(row['reduccion_inercia']) else '-',
                      f"{row['reduccion_porcentual']:.2f}" if pd.notna(row['reduccion_porcentual']) else '-',
                      f"{row['silhouette_score']:.4f}"])
make_table(codo_headers, codo_rows)

doc.add_paragraph(
    'Se observa que la reducción porcentual más significativa ocurre en k=3 (7.22%), mientras que a partir '
    'de k=4 la reducción disminuye progresivamente (4.18%, 3.63%, etc.). Esto sugiere un codo en k=3.'
)

add_figure('fig_codo.png', 'Figura 1: Método del codo - Inercia según número de clusters (k=2 a k=10)')

doc.add_heading('5.2 Coeficiente de Silhouette', level=2)
doc.add_paragraph(
    'El coeficiente de Silhouette mide qué tan similares son los puntos dentro de un cluster comparado '
    'con otros clusters. Su valor oscila entre -1 (mal asignado) y 1 (bien asignado). Los resultados fueron:'
)

sil_rows = []
for _, row in comp.iterrows():
    sil_rows.append([int(row['k']), f"{row['silhouette_score']:.4f}"])
make_table(['k', 'Silhouette Score'], sil_rows)

add_observation(
    'El mayor Silhouette Score se obtuvo con k=2 (0.0986). Sin embargo, la diferencia con k=3 (0.0823) es '
    'pequeña, y k=3 ofrece una interpretación comercial más valiosa (segmentos bajo, medio y alto valor).'
)

add_figure('fig_silhouette.png', 'Figura 2: Silhouette Score según número de clusters (k=2 a k=10)')

doc.add_heading('5.3 Comparación Codo + Silhouette', level=2)
doc.add_paragraph(
    'La tabla comparativa integra ambos criterios:'
)
make_table(codo_headers, codo_rows)

doc.add_paragraph(
    'Análisis de la comparación:\n'
    '- k=2 tiene el mejor Silhouette (0.0986) pero solo dos grupos (bajo-alto), poco granular.\n'
    '- k=3 tiene el segundo mejor Silhouette (0.0823) y el codo más claro (7.22% de reducción).\n'
    '- k=4 a k=10 muestran reducciones decrecientes y Silhouette descendente.'
)

doc.add_heading('5.4 Decisión de k=3', level=2)
doc.add_paragraph(
    'Se seleccionó k=3 por las siguientes razones técnicas y de negocio:\n\n'
    '1. Método del codo: la mayor caída porcentual ocurre en k=3 (7.22%), definiendo el punto de inflexión.\n'
    '2. Silhouette Score: el valor de k=3 (0.0823) es el segundo más alto y muy cercano al máximo (k=2, 0.0986).\n'
    '3. Interpretabilidad: tres segmentos corresponden a perfiles naturales de clientes (menor valor, intermedio, mayor valor).\n'
    '4. Distribución balanceada: los clusters resultantes tienen 161, 106 y 153 clientes, sin grupos extremadamente pequeños.\n'
    '5. Utilidad comercial: tres segmentos permiten diseñar estrategias diferenciadas (retención, crecimiento, fidelización).'
)

add_observation(
    'La selección de k=3 no coincide con el valor de máximo Silhouette (k=2), pero está justificada '
    'por el codo, la interpretabilidad del negocio y la distribución equilibrada de los clusters.'
)

# ===== 6. MODELO FINAL =====
doc.add_heading('6. Modelo Final K-Means (k=3)', level=1)
doc.add_paragraph(
    'Con k=3 seleccionado, se entrenó el modelo final K-Means con random_state=42 y n_init=10. '
    'La distribución de clientes por cluster fue:'
)

cluster_counts = [161, 106, 153]
dist_headers = ['Cluster', 'Cantidad de Clientes', 'Porcentaje']
dist_rows = []
total = 420
for i, c in enumerate(cluster_counts):
    dist_rows.append([str(i), str(c), f"{c/total*100:.1f}%"])
make_table(dist_headers, dist_rows)

doc.add_paragraph(
    'El Cluster 0 agrupa 161 clientes (38.3%), el Cluster 1 concentra 106 clientes (25.2%), '
    'y el Cluster 2 reúne 153 clientes (36.4%). La distribución es relativamente equilibrada, '
    'sin clusters dominantes ni grupos marginales.'
)

# ===== 7. PERFILES DE CLUSTERS =====
doc.add_heading('7. Perfiles de los Clusters', level=1)

doc.add_heading('7.1 Perfil Numérico', level=2)
doc.add_paragraph(
    'Se calcularon los promedios de cada variable numérica por cluster para caracterizar los segmentos:'
)

perf_num = pd.read_csv(os.path.join(RUTA, 'perfil_numerico.csv'))
num_headers = ['Cluster'] + list(perf_num.columns[1:])
num_rows = []
for _, row in perf_num.iterrows():
    num_rows.append([str(int(row['Cluster']))] + [f"{v:.2f}" for v in row[1:]])
make_table(num_headers, num_rows)

doc.add_paragraph(
    'Interpretación de los clusters:\n\n'
    'Cluster 0 (161 clientes - "Clientes Estables"):\n'
    '- Edad promedio: 41.4 años. Ingreso: 2,958. Antigüedad: 27 meses (la más baja).\n'
    '- Cantidad de compras: 11.9 (la más baja). Quejas: 0.61 (baja).\n'
    '- Días desde última compra: 34.6 (reciente). Visitas web: 10 (alta).\n'
    '- Cupones usados: 3.1 (el más alto). Tasa de abandono: 36.7%.\n'
    '- Perfil: clientes relativamente nuevos, compran poco pero usan cupones y visitan el sitio web con frecuencia.\n\n'
    'Cluster 1 (106 clientes - "Clientes en Riesgo"):\n'
    '- Edad promedio: 37.1 años (la más joven). Ingreso: 3,137 (el más alto).\n'
    '- Quejas: 2.82 (la más alta, más del triple que los otros grupos).\n'
    '- Días desde última compra: 67 (el más alto). Antigüedad: 42.3 meses.\n'
    '- Tasa de abandono: 65.1% (la más alta, casi el doble).\n'
    '- Perfil: clientes jóvenes con alto ingreso pero altamente insatisfechos, con muchas quejas y alta probabilidad de abandono.\n\n'
    'Cluster 2 (153 clientes - "Clientes Leales"):\n'
    '- Edad promedio: 42.2 años (la más alta). Antigüedad: 71.7 meses (la más alta).\n'
    '- Cantidad de compras: 21.2 (la más alta). Compras últimos 12M: 5.35.\n'
    '- Quejas: 0.81 (baja). Días desde última compra: 36.4.\n'
    '- Tasa de abandono: 34.0% (la más baja).\n'
    '- Perfil: clientes leales y antiguos, con alta frecuencia de compra, pocas quejas y baja intención de abandono.'
)

doc.add_heading('7.2 Perfil Categórico', level=2)
doc.add_paragraph(
    'Se calculó la moda (categoría más frecuente) de cada variable categórica por cluster:'
)

perf_cat = pd.read_csv(os.path.join(RUTA, 'perfil_categorico.csv'))
cat_headers = ['Cluster'] + list(perf_cat.columns[1:])
cat_rows = []
for _, row in perf_cat.iterrows():
    cat_rows.append([str(int(row['Cluster']))] + [str(v) for v in row[1:]])
make_table(cat_headers, cat_rows)

doc.add_paragraph(
    'Análisis del perfil categórico:\n'
    '- Todos los clusters tienen como ciudad modal Bogotá y zona urbana, lo cual refleja la composición general del dataset.\n'
    '- El canal preferido es Web en los tres clusters, indicando una tendencia general hacia canales digitales.\n'
    '- El segmento modal en todos los clusters es Básico, lo cual es consistente con la distribución del dataset.\n'
    '- La diferencia clave está en Satisfacción: el Cluster 1 (en riesgo) tiene moda Baja, mientras que los Clusters 0 y 2 tienen moda Media.'
)

doc.add_heading('7.3 Análisis de Abandono por Cluster', level=2)
doc.add_paragraph(
    'La variable Abandono se utilizó como validación externa para evaluar la utilidad de los segmentos:'
)

abandono_headers = ['Cluster', 'Cantidad Clientes', 'Tasa de Abandono (%)']
abandono_rows = [
    ['0 - Estables', '161', '36.65'],
    ['1 - En Riesgo', '106', '65.09'],
    ['2 - Leales', '153', '33.99'],
]
make_table(abandono_headers, abandono_rows)

doc.add_paragraph(
    'El análisis de abandono valida los segmentos encontrados:\n'
    '- Cluster 1 (En Riesgo): 65.09% de abandono, casi el doble que los otros grupos.\n'
    '- Cluster 0 (Estables): 36.65% de abandono, nivel medio.\n'
    '- Cluster 2 (Leales): 33.99% de abandono, el más bajo.\n\n'
    'Esto demuestra que los clusters identificados tienen poder discriminante respecto al comportamiento '
    'de abandono, lo cual es valioso para estrategias de retención.'
)

add_figure('fig_barras.png', 'Figura 3: Distribución de clientes por cluster (gráfico de barras)', width=4.5)

# ===== 8. VISUALIZACIONES =====
doc.add_heading('8. Visualizaciones', level=1)

doc.add_heading('8.1 Dispersión sin Clusters', level=2)
doc.add_paragraph(
    'La gráfica de dispersión de Ingreso Mensual vs. Cantidad de Compras sin clusterizar '
    'muestra la distribución general de los datos antes de la segmentación. Se observa '
    'una nube de puntos sin una separación evidente a simple vista.'
)
add_figure('fig_dispersion_sin.png', 'Figura 4: Dispersión inicial - Ingreso mensual vs Cantidad de compras')

doc.add_heading('8.2 Dispersión con Clusters (k=3)', level=2)
doc.add_paragraph(
    'Al colorear los puntos según el cluster asignado, se aprecia una separación parcial '
    'entre los grupos, especialmente en el eje de Cantidad de Compras. El Cluster 2 (Leales) '
    'tiende a concentrarse en valores altos de compras, mientras que el Cluster 1 (En Riesgo) '
    'se distribuye con mayor dispersión.'
)
add_figure('fig_dispersion_con.png', 'Figura 5: Dispersión con clusters - Ingreso mensual vs Cantidad de compras')

doc.add_heading('8.3 Visualización PCA', level=2)
doc.add_paragraph(
    'PCA reduce las 24 dimensiones de la matriz preprocesada a 2 componentes principales, '
    'permitiendo visualizar la estructura general de los clusters.'
)
doc.add_paragraph(
    'Varianza explicada: PC1 = 12.22%, PC2 = 12.09%, total = 24.31%. '
    'Aunque la varianza explicada es moderada (lo cual es esperado con datos heterogéneos), '
    'la proyección visual muestra cierta separación entre los clusters.'
)
add_figure('fig_pca.png', 'Figura 6: Visualización de clusters con PCA (2 componentes)')

doc.add_heading('8.4 Distribución por Cluster', level=2)
doc.add_paragraph(
    'La gráfica de barras muestra la cantidad de clientes en cada cluster: '
    '161 en el Cluster 0, 106 en el Cluster 1 y 153 en el Cluster 2.'
)
# fig_barras already added above

doc.add_heading('8.5 Diagramas de Caja (Boxplots)', level=2)
doc.add_paragraph(
    'Los diagramas de caja permiten comparar la distribución de las variables numéricas '
    'entre los clusters, identificando medianas, cuartiles y valores atípicos.'
)
add_figure('fig_boxplots.png', 'Figura 7: Diagramas de caja por cluster para las principales variables numéricas')

doc.add_heading('8.6 Matriz de Correlación', level=2)
doc.add_paragraph(
    'La matriz de correlación muestra las relaciones lineales entre las variables numéricas. '
    'Se observan correlaciones positivas entre IngresoMensual y CantidadCompras, y entre '
    'AntiguedadMeses y CantidadCompras, mientras que QuejasUltimos6M muestra correlación '
    'negativa con Satisfacción.'
)
add_figure('fig_correlacion.png', 'Figura 8: Matriz de correlación de las variables numéricas')

# ===== 9. CONCLUSIONES =====
doc.add_heading('9. Conclusiones', level=1)

conclusions = [
    'Se segmentó exitosamente la base de 420 clientes en 3 grupos homogéneos utilizando K-Means, '
    'demostrando la aplicabilidad del aprendizaje no supervisado en inteligencia de negocios.',
    'El método del codo evidenció un punto de inflexión en k=3 con una reducción de inercia del 7.22%, '
    'mientras que el Silhouette Score confirmó que k=3 es una opción razonable (0.0823).',
    'Los tres segmentos identificados presentan perfiles diferenciados: Clientes Estables (38.3%, '
    'baja actividad, uso de cupones), Clientes en Riesgo (25.2%, altas quejas, alta probabilidad de '
    'abandono) y Clientes Leales (36.4%, alta antigüedad, alta frecuencia de compra, baja tasa de abandono).',
    'El análisis de abandono validó externamente los clusters: el Cluster 1 presenta una tasa de '
    'abandono del 65.09%, casi el doble que los otros grupos, confirmando que los segmentos '
    'capturan diferencias reales en el comportamiento del cliente.',
    'Las visualizaciones (dispersión, PCA, boxplots, correlación) complementaron el análisis '
    'numérico y facilitaron la interpretación de los segmentos.',
    'La metodología aplicada puede extenderse a futuros análisis incluyendo más variables '
    'o probando otros algoritmos de clustering (DBSCAN, clustering jerárquico).',
]
for i, c in enumerate(conclusions):
    doc.add_paragraph(f"{i+1}. {c}")

# ===== 10. RECOMENDACIONES =====
doc.add_heading('10. Recomendaciones', level=1)

recs = [
    'Para el Cluster 1 (En Riesgo): implementar un programa de retención urgente, '
    'investigar las causas de las quejas (2.82 en promedio), ofrecer incentivos personalizados '
    'y realizar seguimiento proactivo para reducir la tasa de abandono del 65%.',
    'Para el Cluster 2 (Leales): diseñar un programa de fidelización que reconozca su lealtad, '
    'ofrecer beneficios exclusivos y aprovechar su alta frecuencia de compra para campañas de '
    'cross-selling y up-selling.',
    'Para el Cluster 0 (Estables): incentivar el aumento de compras mediante promociones '
    'dirigidas, aprovechar su alta disposición al uso de cupones y fomentar su transición '
    'hacia clientes de mayor valor.',
    'Monitorear periódicamente la asignación de clusters para detectar cambios en el '
    'comportamiento de los clientes y ajustar las estrategias comerciales.',
]
for i, r in enumerate(recs):
    doc.add_paragraph(f"{i+1}. {r}")

# ===== 11. REFERENCIAS =====
doc.add_heading('11. Referencias', level=1)
refs = [
    'SENA. (2026). Guía 4 - Preparación de datasets para aprendizaje no supervisado. Centro de Gestión de Mercados, Logística y Tecnologías de la Información.',
    'Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.',
    'McKinney, W. (2010). Data Structures for Statistical Computing in Python. Proceedings of the 9th Python in Science Conference.',
    'Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. Computing in Science & Engineering, 9(3), 90-95.',
]
for i, r in enumerate(refs):
    doc.add_paragraph(f"{i+1}. {r}")

# ---- Guardar ----
output = os.path.join(RUTA, 'informe final completo.docx')
doc.save(output)
print(f"✅ Informe generado: {output}")

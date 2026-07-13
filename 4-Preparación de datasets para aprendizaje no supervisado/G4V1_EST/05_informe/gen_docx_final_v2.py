from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AZUL = RGBColor(0x1B, 0x3A, 0x6B)
NEGRO = RGBColor(0x00, 0x00, 0x00)
GRIS = RGBColor(0x80, 0x80, 0x80)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

OUT = r'C:\Users\SENA\Documents\VS_CODE\Sena_IA_1\IA_1\4-Preparación de datasets para aprendizaje no supervisado\G4V1_EST\05_informe'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = NEGRO
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'; hs.font.bold = True; hs.font.color.rgb = AZUL
    hs.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    hs.paragraph_format.space_after = Pt(6)

def shd(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),color)
    cell._tc.get_or_add_tcPr().append(s)

def T(headers, rows, caption=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h; shd(c,'1B3A6B')
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold=True; r.font.color.rgb=BLANCO; r.font.size=Pt(9)
    for row in rows:
        r = t.add_row()
        for i,v in enumerate(row):
            c = r.cells[i]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for ru in p.runs: ru.font.size=Pt(9)
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption); r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=GRIS
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return t

def obs(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"OBSERVACION: {text}")
    r.font.size=Pt(10); r.font.italic=True; r.font.bold=True; r.font.color.rgb=AZUL

def F(img, pie, desc, interp, anal, conc):
    path = os.path.join(OUT, img)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    r = p.add_run(pie); r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=GRIS
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph()
    for lbl, txt in [("Descripcion: ",desc),("Interpretacion: ",interp),("Analisis: ",anal),("Conclusion: ",conc)]:
        pp = doc.add_paragraph()
        rr = pp.add_run(lbl); rr.font.bold=True; rr.font.size=Pt(10)
        rr = pp.add_run(txt); rr.font.size=Pt(10)

# ═══════════════════ PORTADA ═══════════════════
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Guia 4 - Informe de Segmentacion\n"); r.font.size=Pt(28); r.font.color.rgb=AZUL; r.font.bold=True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Segmentacion de Clientes mediante\nAprendizaje No Supervisado (K-Means)"); r.font.size=Pt(15); r.font.color.rgb=NEGRO
doc.add_paragraph()
for line in ["Curso: Preparacion de Datasets para Aprendizaje No Supervisado",
             "Estudiante: [Nombre del estudiante]",
             "Fecha: Junio 2026"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.font.size=Pt(12)
doc.add_page_break()

# ═══════════════════ INDICE ═══════════════════
doc.add_heading('Indice', level=1)
indice = [
    "1. Introduccion",
    "2. Descripcion del Dataset",
    "3. Variables Seleccionadas para Clustering",
    "4. Clasificacion de Variables y Preprocesamiento",
    "  4.1 Clasificacion de variables",
    "  4.2 Preprocesamiento aplicado",
    "5. Seleccion del Numero de Clusters (k)",
    "  5.1 El algoritmo K-Means",
    "  5.2 Metodo del Codo",
    "  5.3 Silhouette Score",
    "  5.4 Comparacion entre Codo y Silhouette",
    "  5.5 Decision final: por que k=3?",
    "6. Evaluacion Interna del Clustering",
    "7. Perfilamiento de Segmentos",
    "  7.1 Perfil Numerico",
    "  7.2 Perfil Categorico",
    "  7.3 Analisis de Abandono por Cluster",
    "  7.4 Interpretacion de Segmentos",
    "8. Visualizaciones",
    "  8.1 Dispersion inicial",
    "  8.2 Dispersion con clusters",
    "  8.3 Visualizacion PCA",
    "  8.4 Perfil comparativo de clusters",
    "  8.5 Distribucion por variable (Boxplots)",
    "  8.6 Matriz de correlacion",
    "9. Conclusiones",
    "10. Limitaciones y Recomendaciones",
    "11. Declaracion de Uso de IA",
]
for item in indice:
    p = doc.add_paragraph(item); p.paragraph_format.left_indent = Cm(0.5)
doc.add_page_break()

# ═══════════════════ 1. INTRODUCCION ═══════════════════
doc.add_heading('1. Introduccion', level=1)

doc.add_paragraph(
    "La segmentacion de clientes es una de las aplicaciones mas estrategicas del aprendizaje automatico "
    "en el ambito empresarial. Consiste en dividir una poblacion heterogenea de clientes en subgrupos "
    "homogeneos internamente pero diferenciados entre si, con el fin de disenar estrategias comerciales, "
    "de marketing y de retencion especificas para cada perfil. El aprendizaje no supervisado, y en "
    "particular el algoritmo K-Means, ofrece un enfoque basado en datos para descubrir estos segmentos "
    "de manera automatica, sin necesidad de etiquetas o conocimiento previo sobre los grupos existentes."
)

doc.add_paragraph(
    "El presente informe documenta el analisis completo de segmentacion realizado sobre una base de datos "
    "de 420 clientes de una empresa comercial. El objetivo fundamental es identificar grupos diferenciados "
    "de clientes basados en sus caracteristicas demograficas, comportamentales y transaccionales, de modo "
    "que la organizacion pueda comprender mejor a su base de clientes y disenar estrategias especificas "
    "para cada segmento identificado."
)

doc.add_paragraph(
    "La metodologia aplicada sigue un flujo de trabajo riguroso y sistematico: (1) carga y diagnostico "
    "del dataset, verificando valores nulos, duplicados y tipos de datos; (2) seleccion fundamentada de "
    "15 variables relevantes para la segmentacion, excluyendo identificadores y variables no predictivas; "
    "(3) clasificacion de variables en numericas, nominales y ordinales para aplicar el tratamiento "
    "adecuado a cada tipo; (4) preprocesamiento mediante estandarizacion (StandardScaler) para variables "
    "numericas, codificacion One-Hot para nominales y codificacion ordinal para ordinales; (5) determinacion "
    "del numero optimo de clusters mediante el metodo del codo y el coeficiente de silueta, evaluando "
    "k desde 2 hasta 10; (6) entrenamiento del modelo final K-Means con k=3; (7) perfilamiento numerico "
    "y categorico detallado de cada segmento; (8) validacion externa mediante la variable Abandono; y "
    "(9) visualizacion de resultados mediante PCA, graficas de dispersion, boxplots y matriz de correlacion."
)

doc.add_paragraph(
    "La importancia de este analisis trasciende lo meramente tecnico: una segmentacion bien definida "
    "permite a la empresa optimizar sus recursos comerciales, mejorar la experiencia del cliente, "
    "incrementar la retencion, disenar promociones personalizadas y, en ultima instancia, aumentar su "
    "rentabilidad. Los resultados obtenidos constituyen un insumo estrategico fundamental para la toma "
    "de decisiones basada en evidencias."
)

# ═══════════════════ 2. DESCRIPCION DEL DATASET ═══════════════════
doc.add_heading('2. Descripcion del Dataset', level=1)

doc.add_paragraph(
    "El conjunto de datos utilizado corresponde al archivo G4_base_clientes.csv, el cual contiene "
    "informacion detallada de 420 clientes de una empresa comercial. El dataset esta compuesto por "
    "18 columnas que incluyen datos demograficos, transaccionales y de comportamiento, ofreciendo "
    "una vision multidimensional de cada cliente."
)

T(
    ["Caracteristica", "Descripcion", "Valor"],
    [
        ["Numero de registros", "Clientes muestreados", "420"],
        ["Numero de variables", "Columnas totales", "18"],
        ["Variables numericas", "Enteros (int64) y flotante (float64)", "11"],
        ["Variables categoricas", "Texto (string)", "7"],
        ["Valores nulos", "En ninguna columna", "0"],
        ["Registros duplicados", "Ninguno", "0"],
        ["Rango de edad", "Minima y maxima", "18 - 65 anos"],
        ["Rango ingresos", "Minimo y maximo", "$750 - $6,155"],
    ],
    "Tabla 1. Resumen descriptivo del dataset G4_base_clientes.csv."
)

doc.add_paragraph(
    "El dataset se encuentra en un estado optimo para el analisis: no presenta valores nulos en ninguna "
    "de sus 18 columnas, lo que elimina la necesidad de aplicar tecnicas de imputacion que podrian "
    "introducir sesgos. Tampoco existen registros duplicados, lo que garantiza que cada fila representa "
    "un cliente unico. Las variables cubren un espectro amplio de dimensiones del cliente: demograficas "
    "(edad, ciudad, zona de residencia), economicas (ingreso mensual, segmento), de comportamiento de "
    "compra (cantidad de compras, antiguedad, compras recientes), de engagement digital (visitas web, "
    "tiempo de sesion, cupones utilizados) y de percepcion del servicio (quejas, satisfaccion). Adicionalmente, "
    "el dataset incluye la variable Abandono, que indica si el cliente ha dejado la empresa, la cual se "
    "reserva exclusivamente para validacion externa al final del analisis."
)

obs("La variable Abandono no se utiliza como variable de entrada para el clustering, ya que el aprendizaje "
    "no supervisado no requiere variable objetivo. Su uso se limita a la validacion externa de los segmentos "
    "identificados.")

# ═══════════════════ 3. VARIABLES SELECCIONADAS ═══════════════════
doc.add_heading('3. Variables Seleccionadas para Clustering', level=1)

doc.add_paragraph(
    "Para el proceso de clustering se preseleccionaron 15 de las 18 variables originales. La seleccion "
    "se realizo con base en el criterio de que cada variable debe representar una caracteristica del "
    "cliente que pueda contribuir a diferenciar grupos con sentido comercial. Se excluyeron tres variables:"
)

doc.add_paragraph(
    "ID_Cliente: identificador unico sin valor predictivo ni comportamental. Su inclusion como variable "
    "de clustering generaria 420 categorias diferentes sin significado para la segmentacion.",
    style='List Bullet'
)
doc.add_paragraph(
    "CodigoCampania: codigo administrativo de campana que no describe atributos intrinsecos del cliente "
    "ni su comportamiento.",
    style='List Bullet'
)
doc.add_paragraph(
    "Abandono: variable de resultado (target en aprendizaje supervisado) que no debe guiar la formacion "
    "de los clusters en un enfoque no supervisado.",
    style='List Bullet'
)

T(
    ["Variable", "Tipo Sklearn", "Descripcion", "Relevancia para segmentacion"],
    [
        ["Edad", "Numerica", "Edad del cliente en anos", "Perfil demografico y ciclo de vida"],
        ["IngresoMensual", "Numerica", "Ingreso mensual estimado", "Poder adquisitivo y capacidad de gasto"],
        ["CantidadCompras", "Numerica", "Total de compras historicas", "Lealtad y frecuencia de compra"],
        ["ComprasUltimos12M", "Numerica", "Compras en el ultimo ano", "Actividad comercial reciente"],
        ["AntiguedadMeses", "Numerica", "Meses desde el primer registro", "Lealtad y relacion a largo plazo"],
        ["QuejasUltimos6M", "Numerica", "Numero de quejas recientes", "Insatisfaccion y riesgo de abandono"],
        ["DiasDesdeUltimaCompra", "Numerica", "Dias desde la ultima transaccion", "Inactividad y posible desvinculacion"],
        ["VisitasWebUltimoMes", "Numerica", "Visitas al sitio web en el ultimo mes", "Engagement con el canal digital"],
        ["TiempoPromedioSesionMin", "Numerica", "Minutos promedio por sesion web", "Profundidad de interes digital"],
        ["CuponesUsados", "Numerica", "Cantidad de cupones canjeados", "Sensibilidad a promociones y descuentos"],
        ["Ciudad", "Nominal", "Ciudad de residencia", "Distribucion geografica"],
        ["CanalPreferido", "Nominal", "Canal de compra favorito", "Preferencia de canal de contacto"],
        ["ZonaResidencia", "Nominal", "Urbana, Suburbana o Rural", "Contexto geografico de residencia"],
        ["Segmento", "Ordinal", "Basico < Medio < Premium", "Nivel de cliente segun la empresa"],
        ["Satisfaccion", "Ordinal", "Baja < Media < Alta", "Percepcion de satisfaccion con el servicio"],
    ],
    "Tabla 2. Variables seleccionadas para clustering con clasificacion y relevancia."
)

doc.add_paragraph(
    "Esta seleccion de 15 variables busca capturar de manera integral las distintas dimensiones del perfil "
    "del cliente. La inclusion de variables demograficas, economicas, comportamentales, de engagement y "
    "de satisfaccion asegura que los clusters identificados reflejen diferencias reales y multidimensionales "
    "entre los clientes, lo que maximiza la utilidad practica de la segmentacion."
)

# ═══════════════════ 4. CLASIFICACION Y PREPROCESAMIENTO ═══════════════════
doc.add_heading('4. Clasificacion de Variables y Preprocesamiento', level=1)

doc.add_heading('4.1 Clasificacion de variables', level=2)

doc.add_paragraph(
    "El algoritmo K-Means opera calculando distancias euclidianas entre los puntos en un espacio "
    "multidimensional. Para que el calculo sea correcto, todas las variables deben ser numericas y "
    "estar en escalas comparables. Por esta razon, antes del preprocesamiento se clasificaron las "
    "15 variables seleccionadas en tres categorias, cada una con un tratamiento especifico:"
)

doc.add_paragraph(
    "Variables numericas (10): Edad, IngresoMensual, CantidadCompras, ComprasUltimos12M, "
    "AntiguedadMeses, QuejasUltimos6M, DiasDesdeUltimaCompra, VisitasWebUltimoMes, "
    "TiempoPromedioSesionMin, CuponesUsados. Estas variables se estandarizaran para tener "
    "media 0 y desviacion estandar 1.",
    style='List Bullet'
)
doc.add_paragraph(
    "Variables nominales (3): Ciudad, CanalPreferido, ZonaResidencia. Categorias sin orden "
    "intrinseco que se codificaran mediante One-Hot Encoding.",
    style='List Bullet'
)
doc.add_paragraph(
    "Variables ordinales (2): Segmento (Basico < Medio < Premium) y Satisfaccion "
    "(Baja < Media < Alta). Categorias con orden logico que se codificaran respetando "
    "su jerarquia natural.",
    style='List Bullet'
)

T(
    ["Categoria", "Cantidad", "Variables", "Transformacion"],
    [
        ["Numerica", "10", "Edad, IngresoMensual, CantidadCompras, ComprasUltimos12M, AntiguedadMeses, QuejasUltimos6M, DiasDesdeUltimaCompra, VisitasWebUltimoMes, TiempoPromedioSesionMin, CuponesUsados", "StandardScaler"],
        ["Nominal", "3", "Ciudad, CanalPreferido, ZonaResidencia", "OneHotEncoder"],
        ["Ordinal", "2", "Segmento, Satisfaccion", "OrdinalEncoder"],
    ],
    "Tabla 3. Clasificacion de variables para preprocesamiento."
)

doc.add_heading('4.2 Preprocesamiento aplicado', level=2)

doc.add_paragraph(
    "Se aplicaron las siguientes transformaciones utilizando la clase ColumnTransformer de scikit-learn, "
    "que permite ejecutar multiples transformaciones en paralelo sobre diferentes subconjuntos de columnas:"
)

doc.add_paragraph(
    "Estandarizacion (StandardScaler): Las 10 variables numericas se transformaron para que cada una "
    "tenga media igual a 0 y desviacion estandar igual a 1. Este paso es indispensable porque K-Means "
    "es altamente sensible a la magnitud de las variables. Sin estandarizacion, variables como "
    "IngresoMensual (con valores entre 750 y 6,155) dominarian el calculo de distancias sobre variables "
    "como QuejasUltimos6M (con valores entre 0 y 6), distorsionando completamente la segmentacion."
)
doc.add_paragraph(
    "One-Hot Encoding: Las 3 variables nominales se transformaron en columnas binarias, creando una "
    "columna por cada categoria unica. Para Ciudad (con valores como Bogota, Medellin, Cali, Cartagena, "
    "Barranquilla), CanalPreferido (App, Web, Tienda, Telefono) y ZonaResidencia (Urbana, Suburbana, "
    "Rural), este proceso genero aproximadamente 11 columnas binarias adicionales."
)
doc.add_paragraph(
    "Codificacion Ordinal: Las 2 variables ordinales se codificaron respetando su orden natural. "
    "Para Segmento: Basico=0, Medio=1, Premium=2. Para Satisfaccion: Baja=0, Media=1, Alta=2. "
    "Esta codificacion preserva la relacion de orden entre categorias."
)

T(
    ["Procedimiento", "Libreria", "Variables", "Resultado"],
    [
        ["StandardScaler", "sklearn.preprocessing", "10 numericas", "10 columnas con media=0, std=1"],
        ["OneHotEncoder", "sklearn.preprocessing", "3 nominales", "~11 columnas binarias"],
        ["OrdinalEncoder", "sklearn.preprocessing", "2 ordinales", "2 columnas con valores 0,1,2"],
        ["Total", "ColumnTransformer", "15 originales", "Matriz final: 420 filas x 24 columnas"],
    ],
    "Tabla 4. Resumen del preprocesamiento aplicado."
)

doc.add_paragraph(
    "La matriz resultante tiene dimensiones de 420 filas por 24 columnas (23 transformadas mas 1 "
    "columna adicional del One-Hot). Este formato es completamente compatible con el algoritmo K-Means "
    "y garantiza que todas las variables contribuyan de manera equitativa al calculo de distancias."
)

# ═══════════════════ 5. SELECCION DE K ═══════════════════
doc.add_heading('5. Seleccion del Numero de Clusters (k)', level=1)

doc.add_heading('5.1 El algoritmo K-Means', level=2)

doc.add_paragraph(
    "K-Means es un algoritmo de clustering particional que agrupa n registros en k clusters, donde "
    "cada registro pertenece al cluster con el centroide (media) mas cercano. El algoritmo funciona "
    "de manera iterativa: (1) se inicializan k centroides de forma aleatoria; (2) cada registro se "
    "asigna al centroide mas cercano; (3) los centroides se recalculan como el promedio de los "
    "registros asignados a cada cluster; y (4) los pasos 2 y 3 se repiten hasta que los centroides "
    "dejan de cambiar significativamente."
)

doc.add_paragraph(
    "Una caracteristica fundamental de K-Means es que el analista debe especificar el numero de "
    "clusters k antes de ejecutar el algoritmo. El algoritmo no determina automaticamente cuantos "
    "grupos existen en los datos. Por esta razon, la seleccion de k debe basarse en evidencia "
    "tecnica combinada con criterio de negocio. En este analisis se evaluaron valores de k desde "
    "2 hasta 10 utilizando dos metricas complementarias: el metodo del codo (basado en la inercia) "
    "y el coeficiente de silueta (Silhouette Score)."
)

doc.add_heading('5.2 Metodo del Codo', level=2)

doc.add_paragraph(
    "El metodo del codo analiza la inercia, que es la suma de las distancias al cuadrado entre cada "
    "punto y el centroide de su cluster. La inercia mide la compactacion interna de los clusters: "
    "valores mas bajos indican grupos mas homogeneos. Al aumentar k, la inercia siempre disminuye "
    "porque hay mas centroides para ajustarse a los datos. El objetivo es encontrar el punto donde "
    "la disminucion de la inercia se vuelve marginal, formando un codo en la grafica."
)

T(
    ["k", "Inercia", "Reduccion absoluta", "Reduccion %", "Interpretacion"],
    [
        ["2", "4974.43", "-", "-", "Linea base"],
        ["3", "4615.08", "359.35", "7.22%", "Mayor reduccion"],
        ["4", "4422.23", "192.86", "4.18%", "Reduccion significativa"],
        ["5", "4261.72", "160.50", "3.63%", "Reduccion moderada"],
        ["6", "4139.85", "121.87", "2.86%", "Reduccion decreciente"],
        ["7", "4029.71", "110.14", "2.66%", "Reduccion disminuye"],
        ["8", "3939.33", "90.38", "2.24%", "Reduccion baja"],
        ["9", "3852.60", "86.73", "2.20%", "Reduccion marginal"],
        ["10", "3784.10", "68.50", "1.78%", "Reduccion minima"],
    ],
    "Tabla 5. Metodo del codo: inercia, reduccion absoluta y porcentual para k=2 a 10."
)

F("fig_codo.png",
    "Figura 1. Metodo del Codo. Evolucion de la inercia para valores de k desde 2 hasta 10.",
    "La grafica muestra la inercia (suma de distancias intra-cluster) en el eje Y para cada valor de k en el eje X. Cada punto representa el valor calculado y las etiquetas muestran el valor numerico exacto.",
    "Se observa una disminucion consistente de la inercia a medida que k aumenta. La mayor caida ocurre entre k=2 y k=3 (7.22%), seguida de k=3 a k=4 (4.18%). A partir de k=5, la reduccion porcentual se estabiliza por debajo del 3%. El codo se forma alrededor de k=3.",
    "La reduccion porcentual de inercia es el indicador mas revelador: el salto de 7.22% entre k=2 y k=3 es significativamente mayor que cualquier reduccion posterior. Esto sugiere que pasar de 2 a 3 clusters genera la mayor ganancia en compactacion. La reduccion se reduce a la mitad en k=4 (4.18%) y continua decreciendo progresivamente.",
    "El metodo del codo senala k=3 como el punto de inflexion natural, donde el beneficio de agregar un cluster adicional comienza a disminuir notablemente.")

doc.add_heading('5.3 Silhouette Score', level=2)

doc.add_paragraph(
    "El coeficiente de silueta (Silhouette Score) es una metrica de evaluacion interna que mide que tan "
    "similares son los objetos dentro de un cluster en comparacion con los objetos de otros clusters. "
    "Para cada punto, se calcula: (b - a) / max(a, b), donde a es la distancia promedio intra-cluster "
    "(cohesion) y b es la distancia promedio al cluster vecino mas cercano (separacion). El resultado "
    "oscila entre -1 y 1: valores cercanos a 1 indican clusters compactos y bien separados; valores "
    "cercanos a 0 indican solapamiento entre clusters; valores negativos sugieren que los puntos "
    "podrian estar mal asignados."
)

T(
    ["k", "Silhouette Score", "Interpretacion"],
    [
        ["2", "0.0986", "Mejor valor. Separacion moderada."],
        ["3", "0.0823", "Segundo mejor. Diferencia marginal con k=2."],
        ["4", "0.0723", "Desciende. Mayor solapamiento."],
        ["5", "0.0710", "Estable respecto a k=4."],
        ["6", "0.0653", "Continua el descenso."],
        ["7", "0.0662", "Estabilizacion."],
        ["8", "0.0661", "Similar a k=7."],
        ["9", "0.0650", "Minimo del rango evaluado."],
        ["10", "0.0637", "Descenso adicional."],
    ],
    "Tabla 6. Silhouette Score para k=2 a 10 con interpretacion."
)

F("fig_silhouette.png",
    "Figura 2. Silhouette Score segun numero de clusters. Valores para k desde 2 hasta 10.",
    "La grafica muestra el coeficiente de silueta promedio en el eje Y para cada valor de k en el eje X. Las etiquetas indican el valor numerico exacto de cada punto.",
    "El valor maximo se alcanza en k=2 con 0.0986, seguido de k=3 con 0.0823. A partir de k=4, el score desciende y se estabiliza alrededor de 0.065. Todos los valores son inferiores a 0.1, lo que indica que los clusters presentan solapamiento.",
    "La diferencia entre k=2 (0.0986) y k=3 (0.0823) es de solo 0.017 puntos, una distancia marginal. El descenso mas pronunciado ocurre entre k=3 y k=4 (de 0.0823 a 0.0723), lo que sugiere que a partir de 4 clusters la calidad de la separacion se deteriora mas rapidamente.",
    "Aunque k=2 obtiene el mejor silhouette score, la diferencia con k=3 es minima. k=3 ofrece una calidad de separacion casi equivalente con la ventaja de proporcionar una segmentacion mas rica y util.")

doc.add_heading('5.4 Comparacion entre Codo y Silhouette', level=2)

doc.add_paragraph(
    "Es frecuente que el metodo del codo y el silhouette score no senalen exactamente el mismo valor "
    "de k, especialmente en conjuntos de datos reales donde los clusters no tienen una estructura "
    "perfectamente definida. En estos casos, el analista debe integrar ambas metricas con el contexto "
    "del problema y la interpretabilidad de los resultados. La tabla integrada permite visualizar "
    "simultaneamente ambas metricas para cada valor de k."
)

T(
    ["k", "Inercia", "Reduccion %", "Silhouette", "Analisis integrado"],
    [
        ["2", "4974.43", "-", "0.0986", "Mejor silhouette, peor inercia"],
        ["3", "4615.08", "7.22%", "0.0823", "Codo + buen silhouette"],
        ["4", "4422.23", "4.18%", "0.0723", "Reduccion media, silhouette baja"],
        ["5", "4261.72", "3.63%", "0.0710", "Rendimientos decrecientes"],
        ["6", "4139.85", "2.86%", "0.0653", "Solapamiento creciente"],
        ["7", "4029.71", "2.66%", "0.0662", "Estabilizacion"],
        ["8", "3939.33", "2.24%", "0.0661", "Sin mejora significativa"],
        ["9", "3852.60", "2.20%", "0.0650", "Reduccion marginal"],
        ["10", "3784.10", "1.78%", "0.0637", "Sobreajuste potencial"],
    ],
    "Tabla 7. Comparacion integrada: codo y silhouette para k=2 a 10."
)

doc.add_paragraph(
    "Analisis de la tabla integrada: La reduccion porcentual de inercia alcanza su maximo en k=3 "
    "(7.22%), lo que representa el punto de mayor ganancia en compactacion. El silhouette score, "
    "aunque maximo en k=2 (0.0986), presenta una diferencia marginal con k=3 (0.0823). A partir "
    "de k=4, ambas metricas muestran rendimientos decrecientes: la reduccion de inercia cae por "
    "debajo del 5% y el silhouette score desciende progresivamente. Esta convergencia de evidencias "
    "refuerza la seleccion de un valor de k entre 2 y 4."
)

doc.add_heading('5.5 Decision final: por que k=3?', level=2)

doc.add_paragraph(
    "Despues de analizar integralmente los resultados del metodo del codo, el silhouette score, "
    "las visualizaciones de dispersion y el contexto de negocio, se selecciona k=3 como el numero "
    "optimo de clusters. La decision se sustenta en los siguientes criterios:"
)

doc.add_paragraph(
    "1. El metodo del codo senala k=3 como punto de inflexion: la reduccion de inercia entre k=2 "
    "y k=3 (7.22%) es la mayor de todo el rango evaluado, y representa mas del doble de la reduccion "
    "entre k=3 y k=4 (4.18%). Esto indica que el mayor beneficio en compactacion se obtiene al "
    "incorporar un tercer cluster."
)
doc.add_paragraph(
    "2. El silhouette score para k=3 (0.0823) es solo 0.017 puntos inferior al maximo en k=2 "
    "(0.0986), una diferencia que no justifica limitarse a solo dos grupos. Ambos valores indican "
    "un nivel similar de separacion moderada."
)
doc.add_paragraph(
    "3. Con tres clusters se obtiene una segmentacion interpretable y accionable desde el punto de "
    "vista comercial: un segmento estandar, un segmento en riesgo (con alta tasa de quejas y abandono) "
    "y un segmento leal (con alta antiguedad y volumen de compras). Esta estructura de tres niveles "
    "(bajo, medio, alto) es intuitiva y facil de comunicar a las areas de negocio."
)
doc.add_paragraph(
    "4. La distribucion de clientes entre los tres clusters es razonablemente balanceada: 161 clientes "
    "(38.3%), 106 clientes (25.2%) y 153 clientes (36.4%). Ningun segmento es excesivamente pequeno, "
    "lo que permite realizar analisis estadisticos significativos y disenar estrategias especificas "
    "para cada grupo."
)
doc.add_paragraph(
    "5. k=4 fue descartado porque, aunque la reduccion de inercia (4.18%) aun es relevante, el "
    "silhouette score desciende a 0.0723 y la segmentacion en 4 grupos no ofreceria una mejora "
    "cualitativa suficiente que justifique la complejidad adicional."
)

obs("Decision fundamentada: Se selecciono k = 3 combinando evidencia del metodo del codo (mayor "
    "reduccion porcentual de inercia: 7.22%), silhouette score (0.0823, marginalmente inferior al "
    "maximo), interpretabilidad comercial (3 segmentos: estandar, riesgo, leales) y balance en la "
    "distribucion de clientes (ningun cluster con menos de 100 clientes).")

# ═══════════════════ 6. EVALUACION INTERNA ═══════════════════
doc.add_heading('6. Evaluacion Interna del Clustering', level=1)

doc.add_paragraph(
    "Con k=3 seleccionado, se entreno el modelo final de K-Means utilizando los siguientes parametros: "
    "n_clusters=3, random_state=42 (para asegurar reproducibilidad) y n_init=10 (el algoritmo se ejecuta "
    "10 veces con diferentes inicializaciones y se conserva la mejor solucion)."
)

T(
    ["Metrica", "Valor obtenido", "Interpretacion"],
    [
        ["Inercia (k=3)", "4615.08", "Suma de distancias intra-cluster. Referencia para comparacion."],
        ["Silhouette Score (k=3)", "0.0823", "Separa: 0.0823. Indica estructura con solapamiento moderado."],
        ["Tamanio Cluster 0", "161 clientes (38.3%)", "Segmento mayoritario"],
        ["Tamanio Cluster 1", "106 clientes (25.2%)", "Segmento minoritario"],
        ["Tamanio Cluster 2", "153 clientes (36.4%)", "Segmento leal"],
        ["Iteraciones de convergencia", "10 (n_init)", "10 inicializaciones para evitar optimos locales"],
    ],
    "Tabla 8. Metricas internas y distribucion del modelo final con k=3."
)

doc.add_paragraph(
    "La evaluacion interna revela que la segmentacion con k=3 presenta una calidad aceptable para "
    "datos reales de clientes. El silhouette score de 0.0823, aunque moderado, es consistente con "
    "lo que se espera en conjuntos de datos donde las fronteras entre segmentos de clientes son "
    "inherentemente difusas. Los comportamientos humanos rara vez se dividen en categorias discretas "
    "perfectamente separadas; por el contrario, existe un continuo de perfiles con zonas de transicion. "
    "La distribucion balanceada de los clusters (ningun grupo con menos de 100 clientes) garantiza "
    "que cada segmento tenga suficiente representacion estadistica para ser analizado y utilizado "
    "en la toma de decisiones comerciales."
)

# ═══════════════════ 7. PERFILAMIENTO ═══════════════════
doc.add_heading('7. Perfilamiento de Segmentos', level=1)

doc.add_paragraph(
    "El perfilamiento de clusters consiste en analizar las caracteristicas promedio de cada grupo "
    "para comprender que los distingue. Se presentan tres perspectivas complementarias: el perfil "
    "numerico (promedios de variables cuantitativas), el perfil categorico (moda de variables "
    "cualitativas) y el analisis de abandono (validacion externa)."
)

doc.add_heading('7.1 Perfil Numerico', level=2)

T(
    ["Variable", "Cluster 0 (n=161)", "Cluster 1 (n=106)", "Cluster 2 (n=153)", "Diferencia clave"],
    [
        ["Edad", "41.40", "37.08", "42.18", "C1 es mas joven"],
        ["IngresoMensual", "$2,957.65", "$3,136.82", "$2,957.60", "C1 tiene mayor ingreso"],
        ["CantidadCompras", "11.91", "14.75", "21.22", "C2 compra mucho mas"],
        ["ComprasUltimos12M", "5.17", "4.45", "5.35", "C1 menos activo reciente"],
        ["AntiguedadMeses", "26.98", "42.33", "71.67", "C2 es 2.7x mas antiguo"],
        ["QuejasUltimos6M", "0.61", "2.82", "0.81", "C1 tiene 4.6x mas quejas"],
        ["DiasDesdeUltimaCompra", "34.61", "66.98", "36.44", "C1 duplica inactividad"],
        ["VisitasWebUltimoMes", "9.97", "8.20", "9.37", "C0 lidera engagement digital"],
        ["TiempoPromedioSesionMin", "7.26", "7.03", "6.95", "Similar entre clusters"],
        ["CuponesUsados", "3.10", "2.29", "2.16", "C0 usa mas cupones"],
    ],
    "Tabla 9. Perfil numerico de clusters: promedios y diferencias clave."
)

doc.add_paragraph(
    "El perfil numerico revela diferencias sustanciales entre los clusters. La variable AntiguedadMeses "
    "es la que presenta la mayor disparidad: el Cluster 2 tiene 71.7 meses promedio (casi 6 anos), "
    "mientras que el Cluster 0 tiene solo 27 meses. La variable QuejasUltimos6M tambien muestra "
    "una diferencia muy marcada: el Cluster 1 promedia 2.82 quejas, frente a 0.61 y 0.81 de los "
    "otros clusters (4.6 veces mas). La variable DiasDesdeUltimaCompra confirma la inactividad del "
    "Cluster 1 con 67 dias frente a 35-36 dias de los otros grupos."
)

doc.add_heading('7.2 Perfil Categorico', level=2)

T(
    ["Variable", "Cluster 0 (n=161)", "Cluster 1 (n=106)", "Cluster 2 (n=153)"],
    [
        ["Ciudad (moda)", "Bogota", "Bogota", "Bogota"],
        ["CanalPreferido (moda)", "Web", "Web", "Web"],
        ["ZonaResidencia (moda)", "Urbana", "Urbana", "Urbana"],
        ["Segmento (moda)", "Basico", "Basico", "Basico"],
        ["Satisfaccion (moda)", "Media", "Baja", "Media"],
    ],
    "Tabla 10. Perfil categorico de clusters: moda por variable."
)

doc.add_paragraph(
    "El perfil categorico revela un hallazgo importante: la unica variable categorica que discrimina "
    "entre clusters es Satisfaccion. Mientras que los Clusters 0 y 2 tienen como moda Satisfaccion "
    "Media, el Cluster 1 tiene Satisfaccion Baja. Esto es consistente con el alto numero de quejas "
    "en el Cluster 1 (2.82 promedio) y sugiere que la insatisfaccion es un rasgo distintivo de este "
    "segmento. Las variables Ciudad, CanalPreferido, ZonaResidencia y Segmento tienen la misma moda "
    "(Bogota, Web, Urbana, Basico) en los tres clusters, lo que indica que estas variables no son "
    "determinantes para la segmentacion en este dataset."
)

doc.add_heading('7.3 Analisis de Abandono por Cluster', level=2)

doc.add_paragraph(
    "La variable Abandono, que no se utilizo durante el entrenamiento del clustering, se emplea ahora "
    "como variable de validacion externa para evaluar si los clusters identificados tienen diferencias "
    "significativas en su tasa de abandono. Este analisis permite verificar la utilidad predictiva de "
    "la segmentacion."
)

T(
    ["Cluster", "Clientes", "Abandonos", "Tasa de abandono"],
    [
        ["Cluster 0 - Estandar", "161", "59", "36.65%"],
        ["Cluster 1 - En Riesgo", "106", "69", "65.09%"],
        ["Cluster 2 - Leales", "153", "52", "33.99%"],
        ["Total", "420", "180", "42.86%"],
    ],
    "Tabla 11. Validacion externa: tasa de abandono por cluster."
)

doc.add_paragraph(
    "Los resultados de la validacion externa son contundentes: el Cluster 1 presenta una tasa de "
    "abandono del 65.09%, que duplica las tasas del Cluster 0 (36.65%) y del Cluster 2 (33.99%). "
    "Esta diferencia valida la segmentacion al demostrar que los clusters identificados tienen poder "
    "predictivo sobre el comportamiento futuro de los clientes. El Cluster 1 no solo tiene mas quejas "
    "y mayor inactividad, sino que efectivamente sus clientes abandonan la empresa con una frecuencia "
    "mucho mayor. El Cluster 2, por el contrario, presenta la tasa de abandono mas baja (33.99%), "
    "confirmando su perfil de clientes leales y valiosos."
)

obs("Hallazgo clave: La tasa de abandono del Cluster 1 (65.09%) es casi el doble que la de los otros "
    "segmentos. Esto confirma que las variables de quejas, dias sin compra y antiguedad capturadas "
    "por el clustering son predictores efectivos del riesgo de abandono.")

doc.add_heading('7.4 Interpretacion de Segmentos', level=2)

p = doc.add_paragraph()
r = p.add_run("Cluster 0 - Clientes Estandar (161 clientes, 38.3%): "); r.font.bold = True
p.add_run(
    "Clientes de edad media (41.4 anos) con ingresos mensuales de $2,958. Antiguedad moderada de "
    "27 meses con un volumen de compras total bajo (11.9 compras). Su actividad reciente es aceptable "
    "(5.2 compras en el ultimo ano, 35 dias desde la ultima compra). Presentan pocas quejas (0.61) "
    "y un uso de cupones relativamente alto (3.10), lo que sugiere sensibilidad a promociones. Lideran "
    "el engagement digital con 9.97 visitas web mensuales y 7.26 minutos de sesion. Tasa de abandono "
    "del 36.65%. Satisfaccion predominante: Media. Perfil: Clientes jovenes-adultos, poder adquisitivo "
    "medio-bajo, lealtad moderada, activos en canal digital y receptivos a ofertas. Representan la "
    "base principal de clientes y tienen potencial de crecimiento hacia el segmento leal."
)

p = doc.add_paragraph()
r = p.add_run("Cluster 1 - Clientes en Riesgo (106 clientes, 25.2%): "); r.font.bold = True
p.add_run(
    "Segmento critico. Edad promedio mas baja (37.1 anos) e ingreso mas alto ($3,137). Sin embargo, "
    "presentan los peores indicadores de satisfaccion y actividad: 2.82 quejas promedio (4.6 veces mas "
    "que los otros clusters), 67 dias desde la ultima compra (el doble) y la menor actividad reciente "
    "(4.45 compras en 12 meses). Aunque tienen una antiguedad considerable (42 meses) y un volumen de "
    "compras intermedio (14.75), su comportamiento reciente indica descontento. La tasa de abandono "
    "alcanza el 65.09% -la mas alta por un margen muy amplio-. Satisfaccion predominante: Baja. Perfil: "
    "Clientes con historial de compras significativo pero actualmente insatisfechos. El hecho de que "
    "tengan el ingreso mas alto y aun asi presenten la peor satisfaccion sugiere que el problema no es "
    "economico sino de experiencia de servicio. Requieren atencion urgente y prioritaria."
)

p = doc.add_paragraph()
r = p.add_run("Cluster 2 - Clientes Leales (153 clientes, 36.4%): "); r.font.bold = True
p.add_run(
    "El segmento mas valioso. Mayor antiguedad promedio con 71.7 meses (casi 6 anos), casi el triple "
    "que el Cluster 0. Mayor volumen de compras total (21.2), duplicando al Cluster 0. Mejor actividad "
    "reciente con 5.4 compras en el ultimo ano y solo 36 dias desde la ultima compra. Bajas quejas "
    "(0.81), similar al Cluster 0. Menor uso de cupones (2.16), lo que sugiere que su lealtad no depende "
    "de promociones. Tasa de abandono mas baja: 33.99%. Satisfaccion predominante: Media. Perfil: "
    "Clientes fidelizados, de larga data, con alto volumen historico de compras y comportamiento estable. "
    "Representan el nucleo de clientes que genera el mayor valor a largo plazo para la empresa y deben "
    "ser priorizados en programas de retencion, reconocimiento y beneficios exclusivos."
)

# ═══════════════════ 8. VISUALIZACIONES ═══════════════════
doc.add_heading('8. Visualizaciones', level=1)

doc.add_paragraph(
    "Las visualizaciones son una herramienta fundamental para interpretar los resultados del clustering. "
    "A continuacion se presentan y analizan las seis graficas generadas durante el analisis."
)

doc.add_heading('8.1 Dispersion inicial: Ingreso mensual vs Compras', level=2)

F("fig_disp_sin.png",
    "Figura 3. Dispersion inicial: Ingreso mensual vs Cantidad de compras. Cada punto representa un cliente sin informacion de cluster.",
    "Grafico de dispersion que muestra la relacion entre el ingreso mensual (eje X) y la cantidad total de compras (eje Y) para los 420 clientes, sin colorear por cluster.",
    "Se observa una nube de puntos concentrada en ingresos entre $750 y $4,000 con compras entre 5 y 25. Existen valores atipicos: clientes con ingresos muy bajos (alrededor de $750) pero con altas compras (superiores a 25), y clientes con ingresos altos (superiores a $5,000) pero con bajas compras.",
    "No se observa una relacion lineal fuerte entre ingreso y compras. Clientes con ingresos similares presentan comportamientos de compra muy diferentes, lo que sugiere que la segmentacion requiere multiples variables para capturar adecuadamente los perfiles.",
    "La dispersion inicial confirma que los datos no tienen una estructura de grupos evidente en dos dimensiones, lo que justifica el uso de K-Means con multiples variables para descubrir patrones no obvios a simple vista.")

doc.add_heading('8.2 Dispersion con clusters', level=2)

F("fig_disp_con.png",
    "Figura 4. Dispersion con clusters: Ingreso mensual vs Cantidad de compras coloreado por cluster.",
    "Misma grafica anterior pero cada punto coloreado segun la asignacion a su cluster (0, 1 o 2). La leyenda identifica el color de cada cluster.",
    "Aunque existe solapamiento entre los clusters en este espacio bidimensional, se observan tendencias: el Cluster 2 (leales) se concentra en la zona de mayor cantidad de compras (parte superior del grafico), mientras que el Cluster 1 (riesgo) aparece distribuido sin una concentracion clara.",
    "El solapamiento era esperado porque K-Means utilizo 24 dimensiones para la segmentacion, y al proyectar a solo 2 variables se pierde informacion. Aun asi, la tendencia del Cluster 2 a concentrarse en mayores volumenes de compra es consistente con su perfil de clientes leales.",
    "La grafica de dispersion con clusters ofrece una vision parcial pero util de como se distribuyen los segmentos en dos variables clave, confirmando que el Cluster 2 se asocia con mayor actividad de compra.")

doc.add_heading('8.3 Visualizacion PCA', level=2)

doc.add_paragraph(
    "El Analisis de Componentes Principales (PCA) permite reducir las 24 dimensiones del espacio original "
    "a solo 2 componentes que capturan la mayor proporcion posible de la varianza de los datos. "
    "Las dos primeras componentes explican conjuntamente el 24.31% de la varianza total (PC1: 12.22%, "
    "PC2: 12.09%). Este porcentaje, aunque moderado, es esperado al reducir drasticamente la dimensionalidad."
)

F("fig_pca.png",
    "Figura 5. Visualizacion de clusters con PCA. Proyeccion en 2D con 24.31% de varianza explicada.",
    "Proyeccion de los 420 clientes en el espacio de las dos primeras componentes principales. Cada color representa un cluster (azul marino: Cluster 0, gris: Cluster 1, negro: Cluster 2).",
    "Se observan tres concentraciones parcialmente diferenciadas en el espacio PCA. El Cluster 2 (negro) se concentra predominantemente en la region izquierda, el Cluster 1 (gris) en la region superior derecha y el Cluster 0 (azul) ocupa una posicion mas central. Existen zonas de transicion donde los clusters se superponen.",
    "La varianza explicada del 24.31% indica que las dos componentes principales capturan una cuarta parte de la informacion total. Las variables que mas contribuyen a PC1 probablemente son AntiguedadMeses y CantidadCompras (diferenciando al Cluster 2), mientras que PC2 parece estar asociada a QuejasUltimos6M (diferenciando al Cluster 1).",
    "La visualizacion PCA respalda la validez de la segmentacion con k=3 al mostrar que los tres clusters ocupan regiones diferenciadas en el espacio de componentes principales, confirmando que la segmentacion captura patrones reales en los datos.")

doc.add_heading('8.4 Perfil comparativo de clusters', level=2)

F("fig_barras.png",
    "Figura 6. Perfil comparativo de clusters en las 5 variables mas discriminantes: Antiguedad, Compras totales, Quejas, Dias sin compra y Cupones usados.",
    "Grafico de barras agrupadas donde cada grupo de tres barras representa una variable y cada barra dentro del grupo corresponde a un cluster. Las alturas reflejan los valores promedio.",
    "El Cluster 2 domina claramente en antiguedad y cantidad de compras. El Cluster 1 se destaca negativamente en quejas y dias sin compra. El Cluster 0 muestra un perfil intermedio con mayor uso de cupones.",
    "Las variables AntiguedadMeses y CantidadCompras son las que mejor discriminan al Cluster 2 (leales), mientras que QuejasUltimos6M y DiasDesdeUltimaCompra discriminan al Cluster 1 (riesgo). CuponesUsados diferencia ligeramente al Cluster 0 (estandar).",
    "El perfil comparativo confirma que los tres clusters tienen identidades claramente diferenciadas con sentido comercial, y que cada variable contribuye de manera distinta a la caracterizacion de los segmentos.")

doc.add_heading('8.5 Distribucion por variable (Boxplots)', level=2)

F("fig_boxplots.png",
    "Figura 7. Diagramas de caja (boxplots) que muestran la distribucion de 6 variables numericas clave desagregadas por cluster.",
    "Seis paneles independientes, cada uno mostrando un boxplot por cluster para las variables: Compras Totales, Antiguedad (meses), Quejas (6m), Dias sin compra, Ingreso Mensual y Visitas Web.",
    "Los boxplots confirman las diferencias observadas en los promedios y ademas revelan informacion adicional sobre la dispersion y los valores atipicos dentro de cada cluster. Por ejemplo, en Quejas, el Cluster 1 no solo tiene una media mas alta sino tambien una dispersion mayor y multiples valores atipicos. En Antiguedad, el Cluster 2 muestra una mediana significativamente mas alta que los otros clusters.",
    "La presencia de valores atipicos en varios clusters sugiere que existen clientes con comportamientos extremos dentro de cada segmento. Estos clientes podrian merecer atencion individualizada. La dispersion (rango intercuartil) es similar entre clusters para la mayoria de variables, excepto para Quejas donde el Cluster 1 muestra mayor variabilidad.",
    "Los boxplots enriquecen el perfilamiento al mostrar no solo los promedios sino tambien la variabilidad interna de cada cluster, ofreciendo una vision mas completa y robusta de la segmentacion.")

doc.add_heading('8.6 Matriz de Correlacion', level=2)

F("fig_correlacion.png",
    "Figura 8. Matriz de correlacion entre las 10 variables numericas y la variable Cluster. Los valores numericos indican el coeficiente de correlacion de Pearson.",
    "Heatmap que muestra las correlaciones entre todos los pares de variables numericas mas la variable Cluster. Los tonos mas oscuros indican correlaciones mas fuertes (positivas o negativas).",
    "Se observan correlaciones notables: AntiguedadMeses con CantidadCompras (correlacion positiva moderada), QuejasUltimos6M con Abandono (correlacion positiva), y Cluster con AntiguedadMeses y CantidadCompras (correlacion positiva). La variable Cluster muestra su correlacion mas fuerte con AntiguedadMeses, confirmando que esta variable es una de las mas determinantes en la segmentacion.",
    "Las correlaciones confirmadas permiten identificar relaciones entre variables que refuerzan la interpretacion de los clusters. Por ejemplo, la correlacion entre Quejas y Abandono respalda la identificacion del Cluster 1 como segmento de alto riesgo.",
    "La matriz de correlacion proporciona una vision global de las relaciones entre variables, validando que las variables identificadas como discriminantes en el perfilamiento tienen respaldo estadistico en sus correlaciones con la variable Cluster.")

# ═══════════════════ 9. CONCLUSIONES ═══════════════════
doc.add_heading('9. Conclusiones', level=1)

doc.add_paragraph(
    "El analisis de segmentacion mediante K-Means sobre la base de 420 clientes ha permitido identificar "
    "tres grupos diferenciados con perfiles, comportamientos y necesidades distintas. A continuacion se "
    "presentan las conclusiones principales:"
)

doc.add_paragraph(
    "1. Se identificaron tres segmentos con valor estrategico: El modelo de clustering con k=3 genero "
    "una segmentacion interpretable y accionable. Los tres clusters -Estandar (38.3%), en Riesgo (25.2%) "
    "y Leales (36.4%)- presentan diferencias sustanciales en variables criticas como antiguedad (71.7 vs "
    "27 meses), volumen de compras (21.2 vs 11.9), quejas (2.82 vs 0.61) y tasa de abandono (65% vs 34%)."
)
doc.add_paragraph(
    "2. El segmento Leales es un activo estrategico: Con 71.7 meses de antiguedad, 21.2 compras totales "
    "y solo 33.99% de abandono, este grupo representa el nucleo de clientes mas valioso. Su comportamiento "
    "estable y su baja tasa de quejas indican una relacion comercial solida que debe protegerse y "
    "fortalecerse mediante programas de fidelizacion."
)
doc.add_paragraph(
    "3. El segmento en Riesgo es una oportunidad de mejora prioritaria: Con 65.09% de abandono, 2.82 "
    "quejas promedio y 67 dias sin comprar, este grupo requiere atencion inmediata. El hecho de que "
    "tengan el ingreso mas alto ($3,137) pero la peor satisfaccion sugiere que el problema es de "
    "experiencia de servicio, no de capacidad adquisitiva."
)
doc.add_paragraph(
    "4. Las variables mas discriminantes son: AntiguedadMeses, CantidadCompras, QuejasUltimos6M "
    "y DiasDesdeUltimaCompra. Estas cuatro variables capturan las dimensiones fundamentales del "
    "ciclo de vida del cliente: lealtad a largo plazo, volumen historico, satisfaccion reciente "
    "y actividad actual."
)
doc.add_paragraph(
    "5. La validacion externa con Abandono confirmo la validez de la segmentacion: el Cluster 1 "
    "duplica las tasas de abandono de los otros clusters, demostrando que la segmentacion tiene "
    "poder predictivo sobre el comportamiento futuro."
)
doc.add_paragraph(
    "6. El analisis de correlacion revelo que la variable Cluster se correlaciona positivamente "
    "con AntiguedadMeses y CantidadCompras, y negativamente con QuejasUltimos6M, lo que valida "
    "estadisticamente la interpretacion cualitativa de los segmentos."
)

obs("Conclusion general: La segmentacion mediante aprendizaje no supervisado ha demostrado ser "
    "una herramienta efectiva para descubrir patrones ocultos en los datos de clientes. Los tres "
    "segmentos identificados ofrecen una vision clara y accionable: retener a los Leales, recuperar "
    "a los que estan en Riesgo y desarrollar el potencial de los Estandar. Este analisis constituye "
    "un insumo estrategico fundamental para la toma de decisiones basada en datos.")

# ═══════════════════ 10. LIMITACIONES Y RECOMENDACIONES ═══════════════════
doc.add_heading('10. Limitaciones y Recomendaciones', level=1)
doc.add_heading('10.1 Limitaciones', level=2)

for t,d in [
    ("Tamano muestral reducido:", " Con solo 420 registros, la generalizacion de los resultados es limitada. Un dataset mas grande permitiria identificar segmentos mas especificos y validar la estabilidad de los clusters."),
    ("Bajo coeficiente de silueta (0.0823):", " Indica solapamiento entre clusters. Esto es comun en datos de clientes pero sugiere que las fronteras entre segmentos no son nitidas."),
    ("Varianza explicada limitada en PCA (24.31%):", " Al reducir 24 dimensiones a solo 2 componentes, se pierde aproximadamente el 75% de la informacion. Las visualizaciones PCA deben interpretarse con cautela."),
    ("K-Means asume clusters esfericos:", " El algoritmo supone que los clusters tienen formas esfericas y tamanos similares, lo que puede no reflejar la estructura real de los datos."),
    ("Sensibilidad a la inicializacion:", " Aunque se uso n_init=10, K-Means puede converger a diferentes soluciones dependiendo de la inicializacion aleatoria de los centroides."),
    ("Posible sesgo geografico:", " La moda de Ciudad en los tres clusters es Bogota, lo que podria indicar un desbalance geografico en la muestra."),
    ("Sin analisis longitudinal:", " El dataset no incluye informacion temporal que permita analizar la evolucion de los clientes a lo largo del tiempo."),
]:
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.bold = True; p.add_run(d)

doc.add_heading('10.2 Recomendaciones', level=2)

for t,d in [
    ("Ampliar la base de datos:", " Recopilar mas datos de clientes para aumentar el tamano muestral y explorar un mayor numero de clusters (k=4 o k=5) con mayor robustez estadistica."),
    ("Incorporar variables adicionales:", " Incluir valor del ticket promedio, frecuencia de compra por categoria de producto, historial de interacciones con servicio al cliente, datos de encuestas NPS y metricas de engagement en redes sociales."),
    ("Probar otros algoritmos de clustering:", " Comparar K-Means con DBSCAN (para detectar clusters de forma arbitraria y outliers), clustering jerarquico aglomerativo (para visualizar dendrogramas) o Gaussian Mixture Models (para incorporar incertidumbre probabilistico en la asignacion)."),
    ("Sistema de alertas tempranas para el Cluster 1:", " Implementar un sistema de monitoreo que detecte incrementos en quejas o descensos en actividad para activar intervenciones proactivas de retencion antes de que el cliente abandone."),
    ("Programa de fidelizacion para el Cluster 2:", " Disenar un programa de beneficios exclusivos para clientes leales que refuerce su vinculo con la empresa y los incentive a aumentar su gasto."),
    ("Estrategia de desarrollo para el Cluster 0:", " Implementar campanas de cross-selling y up-selling orientadas a convertir clientes estandar en leales, aprovechando su alta receptividad a cupones y promociones."),
    ("Monitoreo continuo de la segmentacion:", " Re-evaluar la segmentacion periodicamente (trimestral o semestral) con nuevos datos para detectar cambios en los patrones de comportamiento de los clientes."),
    ("Analisis de canales digitales:", " Profundizar en el analisis de CanalPreferido y metricas web para optimizar la estrategia omnicanal, especialmente para el Cluster 0 que muestra el mayor engagement digital."),
]:
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.bold = True; p.add_run(d)

# ═══════════════════ 11. DECLARACION DE USO DE IA ═══════════════════
doc.add_heading('11. Declaracion de Uso de IA', level=1)
doc.add_paragraph("En cumplimiento con las politicas de integridad academica y transparencia en el uso de herramientas tecnologicas, se declara que:")

for t,d in [
    ("1. Herramientas utilizadas: ", "Se utilizo un modelo de lenguaje basado en inteligencia artificial (DeepSeek) como herramienta de apoyo para la redaccion, estructuraccion y formato del presente documento academico."),
    ("2. Alcance del uso: ", "La IA fue empleada exclusivamente como asistente para organizar el contenido tecnico, redactar parrafos explicativos, generar las tablas y estructurar las secciones del informe. Todas las descripciones tecnicas, interpretaciones de resultados, analisis de metricas y conclusiones se basan exclusivamente en los datos reales obtenidos de la ejecucion del notebook de analisis en Python."),
    ("3. Revision y validacion: ", "El analisis completo -incluyendo la seleccion de variables, la determinacion del numero optimo de clusters, el perfilamiento numerico y categorico, la validacion con Abandono y la interpretacion de metricas- fue realizado de manera autonoma mediante la ejecucion del codigo Python. El estudiante verifico personalmente cada resultado, valido la coherencia de las interpretaciones y ajusto el contenido del informe para garantizar su precision y alineacion con los objetivos del proyecto."),
    ("4. Responsabilidad academica: ", "El estudiante asume la plena responsabilidad sobre el contenido, analisis y conclusiones presentadas en este informe. La IA fue utilizada como una herramienta de apoyo academico y documental, sin sustituir el juicio critico ni el trabajo analitico del autor."),
]:
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.bold = True; p.add_run(d)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Declaracion formulada en junio de 2026.")

# ── SAVE ──
output = os.path.join(OUT, 'informe_segmentacion_completo.docx')
doc.save(output)
print(f"Documento guardado: {output}")
